"""
Ferramentas de arquivos do Jarvis - v0.5.

O cerebro (Claude) usa estas funcoes para listar, procurar, criar, ler,
mover e apagar arquivos por voz, e para analisar planilhas CSV.

Seguranca (o poder cresce junto com a trava):
- TODA operacao e restrita as pastas do usuario (Desktop, Documentos,
  Downloads, Imagens, Musicas, Videos - resolvidas pela API de pastas
  conhecidas do Windows, entao respeitam o redirecionamento do OneDrive)
  mais as "pastas_extras" do config.json. Fora delas: acesso negado.
- Apagar manda para a LIXEIRA (SHFileOperationW com FOF_ALLOWUNDO),
  nunca apaga definitivo.
- 'apagar_arquivo' e 'mover_arquivo' sao FERRAMENTAS_DESTRUTIVAS no
  cerebro: so executam depois do "Confirma, senhor?" -> "sim" por voz.

Caminhos falados: o Claude passa "desktop/compras.txt", "documentos",
"downloads/fotos" ou um caminho completo; _resolver() traduz e valida.
"""

import csv
import io
import json
import os
import re
import shutil
import unicodedata
from collections import Counter
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent
CONFIG_PATH = BASE_DIR / "config.json"

# Pastas conhecidas do Windows (GUIDs oficiais) - a API devolve o caminho
# REAL, ja considerando o redirecionamento do OneDrive (ex.: o Desktop
# deste usuario vive em C:/Users/vinic/OneDrive/Desktop).
_FOLDERIDS = {
    "desktop":    "{B4BFCC3A-DB2C-424C-B029-7FE99A87C641}",
    "documentos": "{FDD39AD0-238F-46AF-ADB4-6C85480369C7}",
    "downloads":  "{374DE290-123F-4565-9164-39C4925E467B}",
    "imagens":    "{33E28130-4E1E-4676-835A-98395C3BC3BB}",
    "musicas":    "{4BD8D571-6D19-48D3-BE97-422220080E43}",
    "videos":     "{18989B1D-99B5-455B-841C-AB7C74E4DDFC}",
}

# Nomes falados -> chave em _FOLDERIDS
_APELIDOS = {
    "area de trabalho": "desktop",
    "desktop": "desktop",
    "documentos": "documentos",
    "meus documentos": "documentos",
    "documents": "documentos",
    "downloads": "downloads",
    "transferencias": "downloads",
    "imagens": "imagens",
    "fotos": "imagens",
    "pictures": "imagens",
    "musicas": "musicas",
    "musica": "musicas",
    "music": "musicas",
    "videos": "videos",
    "video": "videos",
}

# Os testes apontam para pastas temporarias trocando este dict
# ({nome: Path}); None = usar as pastas reais do Windows.
RAIZES_FORCADAS = None

_cache_raizes = None


def _normalizar(texto):
    texto = unicodedata.normalize("NFD", str(texto).lower())
    return "".join(c for c in texto if unicodedata.category(c) != "Mn").strip()


def _pasta_conhecida(guid_str):
    """Caminho real de uma pasta conhecida (SHGetKnownFolderPath)."""
    import ctypes
    from ctypes import wintypes

    class _GUID(ctypes.Structure):
        _fields_ = [("Data1", wintypes.DWORD), ("Data2", wintypes.WORD),
                    ("Data3", wintypes.WORD), ("Data4", ctypes.c_ubyte * 8)]

    guid = _GUID()
    if ctypes.windll.ole32.CLSIDFromString(ctypes.c_wchar_p(guid_str),
                                           ctypes.byref(guid)) != 0:
        return None
    caminho = ctypes.c_wchar_p()
    if ctypes.windll.shell32.SHGetKnownFolderPath(
            ctypes.byref(guid), 0, None, ctypes.byref(caminho)) != 0:
        return None
    p = Path(caminho.value)
    ctypes.windll.ole32.CoTaskMemFree(caminho)
    return p


def _raizes():
    """{nome: Path} das pastas onde o Jarvis pode mexer."""
    global _cache_raizes
    if RAIZES_FORCADAS is not None:
        return RAIZES_FORCADAS
    if _cache_raizes is None:
        raizes = {}
        for nome, fid in _FOLDERIDS.items():
            try:
                p = _pasta_conhecida(fid)
            except Exception:
                p = None
            if p is None:  # fallback: convencao classica
                p = Path.home() / nome.capitalize()
            raizes[nome] = p
        _cache_raizes = raizes
    raizes = dict(_cache_raizes)
    # pastas extras do config (ex.: pasta de projetos fora do perfil)
    try:
        with open(CONFIG_PATH, encoding="utf-8") as f:
            extras = json.load(f).get("arquivos", {}).get("pastas_extras", [])
        for extra in extras:
            p = Path(os.path.expandvars(str(extra))).expanduser()
            if p.is_dir():
                raizes[_normalizar(p.name)] = p
    except Exception:
        pass
    return raizes


def _permitido(p):
    p = Path(p).resolve()
    for raiz in _raizes().values():
        try:
            p.relative_to(Path(raiz).resolve())
            return True
        except ValueError:
            continue
    return False


def _resolver(caminho_falado, deve_existir=False):
    """Traduz 'desktop/compras.txt' (ou caminho completo) num Path validado.

    Levanta ValueError com mensagem em portugues - o texto vai direto
    para o tool_result e o Claude explica ao usuario.
    """
    txt = str(caminho_falado).strip().strip('"').strip("'")
    if not txt:
        raise ValueError("caminho vazio.")
    p = Path(os.path.expandvars(txt)).expanduser()
    if not p.is_absolute():
        partes = [x for x in re.split(r"[\\/]+", txt) if x.strip()]
        raizes = _raizes()
        chave = _APELIDOS.get(_normalizar(partes[0]), _normalizar(partes[0]))
        base = raizes.get(chave)
        if base is None:
            nomes = ", ".join(sorted(raizes))
            raise ValueError(
                f"pasta '{partes[0]}' desconhecida. Use uma destas: {nomes}; "
                "ou um caminho completo.")
        p = Path(base).joinpath(*partes[1:])
    p = p.resolve()
    if not _permitido(p):
        raise ValueError(
            f"acesso negado: '{p}' esta fora das pastas do usuario "
            "(desktop, documentos, downloads, imagens, musicas, videos).")
    if deve_existir and not p.exists():
        raise ValueError(f"'{p}' nao existe.")
    return p


def _tamanho_legivel(n):
    for unidade in ("B", "KB", "MB", "GB"):
        if n < 1024 or unidade == "GB":
            return f"{n:.0f} {unidade}" if unidade == "B" else f"{n:.1f} {unidade}"
        n /= 1024.0


# ---------- ferramentas ----------

def listar_pasta(pasta):
    try:
        p = _resolver(pasta, deve_existir=True)
    except ValueError as e:
        return f"ERRO: {e}"
    if not p.is_dir():
        return f"ERRO: '{p}' e um arquivo, nao uma pasta."
    try:
        itens = sorted(p.iterdir(),
                       key=lambda x: (x.is_file(), _normalizar(x.name)))
    except OSError as e:
        return f"ERRO ao listar '{p}': {e}"
    linhas = []
    for item in itens[:60]:
        if item.is_dir():
            linhas.append(f"[pasta] {item.name}")
        else:
            try:
                linhas.append(f"{item.name} ({_tamanho_legivel(item.stat().st_size)})")
            except OSError:
                linhas.append(item.name)
    extra = f" (mostrando 60 de {len(itens)})" if len(itens) > 60 else ""
    if not linhas:
        return f"A pasta '{p}' esta vazia."
    return f"Conteudo de '{p}' ({len(itens)} itens){extra}:\n" + "\n".join(linhas)


def procurar_arquivo(nome, pasta=""):
    """Procura arquivos/pastas cujo nome CONTENHA `nome` (sem acentos)."""
    alvo = _normalizar(nome)
    if not alvo:
        return "ERRO: informe o nome (ou parte) do arquivo."
    if pasta:
        try:
            bases = [_resolver(pasta, deve_existir=True)]
        except ValueError as e:
            return f"ERRO: {e}"
    else:
        bases = list(_raizes().values())

    achados, examinados = [], 0
    for base in bases:
        if not Path(base).is_dir():
            continue
        for raiz, pastas, arquivos in os.walk(base):
            # nao desce em pastas ocultas/de sistema (node_modules explode)
            pastas[:] = [d for d in pastas
                         if not d.startswith((".", "__", "$"))
                         and _normalizar(d) != "node_modules"]
            for item in pastas + arquivos:
                examinados += 1
                if alvo in _normalizar(item):
                    achados.append(str(Path(raiz) / item))
                    if len(achados) >= 15:
                        break
            if len(achados) >= 15 or examinados > 20000:
                break
        if len(achados) >= 15:
            break
    if not achados:
        return f"Nenhum arquivo com '{nome}' encontrado."
    return f"Encontrei {len(achados)} resultado(s) para '{nome}':\n" + "\n".join(achados)


def criar_pasta(caminho):
    try:
        p = _resolver(caminho)
    except ValueError as e:
        return f"ERRO: {e}"
    if p.is_dir():
        return f"A pasta '{p}' ja existia."
    try:
        p.mkdir(parents=True)
        return f"Pasta criada: {p}"
    except OSError as e:
        return f"ERRO ao criar a pasta: {e}"


def criar_arquivo(caminho, conteudo):
    try:
        p = _resolver(caminho)
    except ValueError as e:
        return f"ERRO: {e}"
    if p.exists():
        return (f"ERRO: '{p}' ja existe. Para substituir, o usuario precisa "
                "pedir para apagar o antigo antes (com confirmacao).")
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(conteudo, encoding="utf-8")
        return f"Arquivo criado: {p} ({len(conteudo)} caracteres)"
    except OSError as e:
        return f"ERRO ao criar o arquivo: {e}"


def _ler_texto(p, max_bytes=200_000):
    bruto = p.read_bytes()[:max_bytes]
    try:
        return bruto.decode("utf-8-sig")
    except UnicodeDecodeError:
        return bruto.decode("cp1252", errors="replace")


def _ler_pdf(p, max_chars=3000):
    """Texto de um PDF (v0.6, pypdf); truncado."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return "ERRO: leitura de PDF indisponivel (instale o pacote pypdf)."
    try:
        leitor = PdfReader(str(p))
        if leitor.is_encrypted:
            try:
                leitor.decrypt("")
            except Exception:
                return f"ERRO: '{p.name}' e um PDF protegido por senha."
        total_paginas = len(leitor.pages)
        partes, tamanho = [], 0
        for i, pagina in enumerate(leitor.pages):
            texto = (pagina.extract_text() or "").strip()
            if texto:
                partes.append(texto)
                tamanho += len(texto)
            if tamanho >= max_chars or i >= 49:
                break
        texto = "\n".join(partes).strip()
    except Exception as e:
        return f"ERRO ao ler o PDF: {e}"
    if not texto:
        return (f"O PDF '{p.name}' nao tem texto extraivel "
                "(provavelmente e escaneado, so imagem).")
    aviso = (f" (inicio; {total_paginas} paginas no total)"
             if tamanho >= max_chars else f" ({total_paginas} pagina(s))")
    return f"Conteudo de '{p.name}'{aviso}:\n{texto[:max_chars]}"


def _ler_docx(p, max_chars=3000):
    """Texto de um documento Word .docx (v0.6, python-docx); truncado."""
    try:
        import docx
    except ImportError:
        return ("ERRO: leitura de .docx indisponivel "
                "(instale o pacote python-docx).")
    try:
        d = docx.Document(str(p))
        texto = "\n".join(par.text for par in d.paragraphs if par.text.strip())
    except Exception as e:
        return f"ERRO ao ler o documento: {e}"
    if not texto.strip():
        return f"O documento '{p.name}' nao tem texto."
    if len(texto) > max_chars:
        return f"Inicio de '{p.name}' (documento grande):\n{texto[:max_chars]}"
    return f"Conteudo de '{p.name}':\n{texto}"


def _ler_xlsx(p):
    """Espiada numa planilha Excel (v0.6, openpyxl): abas + primeiras linhas.
    Para estatisticas, o caminho certo e o analisar_dados."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ("ERRO: leitura de planilha Excel indisponivel "
                "(instale o pacote openpyxl).")
    try:
        wb = load_workbook(str(p), read_only=True, data_only=True)
        abas = ", ".join(wb.sheetnames)
        ws = wb.active
        linhas = []
        for row in ws.iter_rows(values_only=True):
            linhas.append(" | ".join("" if v is None else str(v) for v in row))
            if len(linhas) >= 20:
                break
        wb.close()
    except Exception as e:
        return f"ERRO ao ler a planilha: {e}"
    corpo = "\n".join(linhas) if linhas else "(planilha vazia)"
    return (f"Planilha '{p.name}' (abas: {abas}; ate 20 linhas da aba ativa - "
            f"use analisar_dados para estatisticas):\n{corpo}")[:3000]


def ler_arquivo(caminho):
    try:
        p = _resolver(caminho, deve_existir=True)
    except ValueError as e:
        return f"ERRO: {e}"
    if p.is_dir():
        return f"ERRO: '{p}' e uma pasta; use listar_pasta."
    sufixo = p.suffix.lower()
    # v0.6 - formatos de documento
    if sufixo == ".pdf":
        return _ler_pdf(p)
    if sufixo == ".docx":
        return _ler_docx(p)
    if sufixo in (".xlsx", ".xlsm"):
        return _ler_xlsx(p)
    if sufixo in (".exe", ".dll", ".zip", ".rar", ".7z", ".png",
                  ".jpg", ".jpeg", ".gif", ".mp3", ".mp4", ".wav",
                  ".pptx", ".doc", ".xls"):
        return (f"ERRO: '{p.name}' nao e um arquivo de texto simples "
                f"({p.suffix}); nao consigo ler o conteudo.")
    try:
        texto = _ler_texto(p)
    except OSError as e:
        return f"ERRO ao ler: {e}"
    if len(texto) > 3000:
        return (f"Inicio de '{p.name}' (arquivo grande, primeiros 3000 "
                f"caracteres):\n{texto[:3000]}")
    return f"Conteudo de '{p.name}':\n{texto}" if texto.strip() else \
        f"O arquivo '{p.name}' esta vazio."


# extensoes em que da para acrescentar/substituir texto (v0.7)
_EXTENSOES_TEXTO = (".txt", ".md", ".csv", ".tsv", ".log", ".json", ".ini",
                    ".cfg", ".yaml", ".yml", "")
_MAX_EDICAO = 2_000_000  # nao edita arquivos maiores que isto (2 MB)


def acrescentar_ao_arquivo(caminho, conteudo):
    """Acrescenta texto ao FIM de um arquivo existente, sem apagar nada.

    Arquivos de texto (lista de compras, notas...) e .docx (paragrafos
    novos). E a ferramenta SEGURA de escrita - nao precisa de confirmacao
    porque nao ha perda possivel de conteudo.
    """
    try:
        p = _resolver(caminho, deve_existir=True)
    except ValueError as e:
        return f"ERRO: {e}"
    if p.is_dir():
        return f"ERRO: '{p}' e uma pasta."
    conteudo = str(conteudo)
    if not conteudo.strip():
        return "ERRO: conteudo vazio."
    sufixo = p.suffix.lower()
    if sufixo == ".docx":
        try:
            import docx
        except ImportError:
            return ("ERRO: edicao de .docx indisponivel "
                    "(instale o pacote python-docx).")
        try:
            d = docx.Document(str(p))
            for linha in conteudo.splitlines() or [conteudo]:
                d.add_paragraph(linha)
            d.save(str(p))
        except Exception as e:
            return f"ERRO ao editar o documento: {e}"
        return f"Acrescentado ao fim do documento '{p.name}'."
    if sufixo not in _EXTENSOES_TEXTO:
        return (f"ERRO: nao sei acrescentar texto em {sufixo} (so arquivos "
                "de texto e .docx; para planilha use escrever_celula).")
    if p.stat().st_size > _MAX_EDICAO:
        return f"ERRO: '{p.name}' e grande demais para editar por voz."
    try:
        # \r\n vira \n antes de reescrever, senao o write_text traduz o
        # \n de novo e as quebras de linha dobram (\r\r\n)
        novo = _ler_texto(p, max_bytes=_MAX_EDICAO).replace("\r\n", "\n")
        if novo and not novo.endswith("\n"):
            novo += "\n"
        novo += conteudo.rstrip("\n") + "\n"
        p.write_text(novo, encoding="utf-8")
    except OSError as e:
        return f"ERRO ao editar: {e}"
    return f"Acrescentado ao fim de '{p.name}' ({len(conteudo)} caracteres)."


def substituir_no_arquivo(caminho, trecho, novo):
    """Substitui um trecho EXATO num arquivo de texto.

    DESTRUTIVA no cerebro (o conteudo antigo se perde): so executa depois
    do "Confirma, senhor?" -> "sim" por voz.
    """
    try:
        p = _resolver(caminho, deve_existir=True)
    except ValueError as e:
        return f"ERRO: {e}"
    if p.is_dir():
        return f"ERRO: '{p}' e uma pasta."
    if p.suffix.lower() not in _EXTENSOES_TEXTO:
        return (f"ERRO: so substituo texto em arquivos de texto; "
                f"'{p.name}' e {p.suffix}.")
    if p.stat().st_size > _MAX_EDICAO:
        return f"ERRO: '{p.name}' e grande demais para editar por voz."
    trecho = str(trecho)
    if not trecho:
        return "ERRO: informe o trecho a substituir."
    try:
        # \r\n vira \n: evita quebras dobradas na reescrita (ver acima)
        texto = _ler_texto(p, max_bytes=_MAX_EDICAO).replace("\r\n", "\n")
    except OSError as e:
        return f"ERRO ao ler: {e}"
    vezes = texto.count(trecho)
    if vezes == 0:
        return (f"ERRO: o trecho '{trecho}' nao existe em '{p.name}'. "
                "Leia o arquivo antes (ler_arquivo) e use o texto exato.")
    try:
        p.write_text(texto.replace(trecho, str(novo)), encoding="utf-8")
    except OSError as e:
        return f"ERRO ao gravar: {e}"
    return f"Substituido em '{p.name}' ({vezes} ocorrencia(s))."


def escrever_celula(caminho, celula, valor):
    """Escreve um valor numa celula de planilha .xlsx (ex.: 'B3').

    DESTRUTIVA no cerebro (sobrescreve o valor antigo): so executa depois
    da confirmacao por voz. Numeros viram numeros de verdade, para as
    formulas da planilha continuarem funcionando.
    """
    try:
        p = _resolver(caminho, deve_existir=True)
    except ValueError as e:
        return f"ERRO: {e}"
    if p.suffix.lower() not in (".xlsx", ".xlsm"):
        return f"ERRO: so escrevo em planilhas .xlsx; '{p.name}' e {p.suffix}."
    celula = str(celula).strip().upper().replace(" ", "")
    if not re.fullmatch(r"[A-Z]{1,3}[1-9][0-9]{0,6}", celula):
        return f"ERRO: '{celula}' nao e uma celula valida (ex.: B3, AA10)."
    try:
        from openpyxl import load_workbook
    except ImportError:
        return ("ERRO: edicao de Excel indisponivel "
                "(instale o pacote openpyxl).")
    try:
        wb = load_workbook(str(p))
        ws = wb.active
        antigo = ws[celula].value
        numero = _para_numero(valor)
        if numero is not None:
            ws[celula] = int(numero) if numero == int(numero) else numero
        else:
            ws[celula] = str(valor)
        wb.save(str(p))
        wb.close()
    except PermissionError:
        return (f"ERRO: '{p.name}' esta aberta em outro programa "
                "(feche o Excel e tente de novo).")
    except Exception as e:
        return f"ERRO ao escrever na planilha: {e}"
    antes = "vazia" if antigo is None else f"'{antigo}'"
    return f"Celula {celula} de '{p.name}': era {antes}, agora '{valor}'."


def mover_arquivo(origem, destino):
    """Move ou renomeia. Se `destino` for uma pasta existente, move para
    dentro dela mantendo o nome."""
    try:
        o = _resolver(origem, deve_existir=True)
        d = _resolver(destino)
    except ValueError as e:
        return f"ERRO: {e}"
    if d.is_dir():
        d = d / o.name
    if d.exists():
        return (f"ERRO: ja existe '{d}'. Nao vou sobrescrever - o usuario "
                "precisa apagar ou renomear o existente antes.")
    try:
        d.parent.mkdir(parents=True, exist_ok=True)
        shutil.move(str(o), str(d))
        return f"Movido: '{o}' -> '{d}'"
    except OSError as e:
        return f"ERRO ao mover: {e}"


def _para_lixeira(p):
    """Manda arquivo/pasta para a Lixeira do Windows (recuperavel)."""
    import ctypes
    from ctypes import wintypes

    class SHFILEOPSTRUCTW(ctypes.Structure):
        _fields_ = [
            ("hwnd", wintypes.HWND),
            ("wFunc", ctypes.c_uint),
            ("pFrom", ctypes.c_wchar_p),
            ("pTo", ctypes.c_wchar_p),
            ("fFlags", ctypes.c_ushort),
            ("fAnyOperationsAborted", wintypes.BOOL),
            ("hNameMappings", ctypes.c_void_p),
            ("lpszProgressTitle", ctypes.c_wchar_p),
        ]

    FO_DELETE = 3
    FOF_ALLOWUNDO, FOF_NOCONFIRMATION = 0x40, 0x10
    FOF_SILENT, FOF_NOERRORUI = 0x4, 0x400

    op = SHFILEOPSTRUCTW()
    op.wFunc = FO_DELETE
    op.pFrom = str(p) + "\0"  # a API exige terminacao dupla em nulo
    op.fFlags = (FOF_ALLOWUNDO | FOF_NOCONFIRMATION
                 | FOF_SILENT | FOF_NOERRORUI)
    res = ctypes.windll.shell32.SHFileOperationW(ctypes.byref(op))
    return res == 0 and not op.fAnyOperationsAborted


def apagar_arquivo(caminho):
    try:
        p = _resolver(caminho, deve_existir=True)
    except ValueError as e:
        return f"ERRO: {e}"
    tipo = "pasta" if p.is_dir() else "arquivo"
    try:
        ok = apagar_arquivo.para_lixeira(p)
    except Exception as e:
        return f"ERRO ao apagar: {e}"
    if ok:
        return f"O {tipo} '{p.name}' foi para a Lixeira (da para recuperar por la)."
    return f"ERRO: o Windows recusou mandar '{p}' para a Lixeira."


# indirecao para os testes nao encherem a Lixeira de verdade
apagar_arquivo.para_lixeira = _para_lixeira


def _para_numero(texto):
    """float de '1.234,56' ou '1234.56'; None se nao for numero."""
    t = str(texto).strip().replace("R$", "").replace("%", "").strip()
    if not t:
        return None
    if "," in t and "." in t:
        t = t.replace(".", "").replace(",", ".")
    elif "," in t:
        t = t.replace(",", ".")
    try:
        return float(t)
    except ValueError:
        return None


def _tabela_de_xlsx(p):
    """Linhas (listas de str) da aba ativa de um .xlsx (openpyxl)."""
    from openpyxl import load_workbook
    wb = load_workbook(str(p), read_only=True, data_only=True)
    ws = wb.active
    linhas = []
    for row in ws.iter_rows(values_only=True):
        valores = ["" if v is None else str(v) for v in row]
        if any(v.strip() for v in valores):
            linhas.append(valores)
        if len(linhas) >= 5001:  # cabecalho + 5000 (mesmo teto do CSV)
            break
    wb.close()
    return linhas


def analisar_dados(caminho):
    """Resumo de uma tabela (CSV/TSV ou Excel .xlsx): linhas, colunas,
    estatisticas por coluna."""
    try:
        p = _resolver(caminho, deve_existir=True)
    except ValueError as e:
        return f"ERRO: {e}"
    sufixo = p.suffix.lower()
    if sufixo in (".xlsx", ".xlsm"):
        # v0.6 - planilha Excel
        try:
            linhas = _tabela_de_xlsx(p)
        except ImportError:
            return ("ERRO: analise de Excel indisponivel "
                    "(instale o pacote openpyxl).")
        except Exception as e:
            return f"ERRO ao abrir a planilha: {e}"
        if not linhas:
            return f"A planilha '{p.name}' esta vazia."
    elif sufixo in (".csv", ".txt", ".tsv"):
        try:
            texto = _ler_texto(p, max_bytes=2_000_000)
        except OSError as e:
            return f"ERRO ao ler: {e}"
        if not texto.strip():
            return f"O arquivo '{p.name}' esta vazio."
        primeira = texto.splitlines()[0]
        delim = max((";", ",", "\t"), key=primeira.count)
        leitor = csv.reader(io.StringIO(texto), delimiter=delim)
        linhas = [l for l in leitor if any(c.strip() for c in l)]
    else:
        return (f"ERRO: sei analisar CSV/TSV e Excel (.xlsx); "
                f"'{p.name}' e {p.suffix}.")
    if len(linhas) < 2:
        return f"ERRO: '{p.name}' nao parece uma tabela (menos de 2 linhas)."

    cabecalho = [c.strip() for c in linhas[0]]
    dados = linhas[1:5001]
    partes = [f"'{p.name}': {len(dados)} linhas"
              + (" (analisei as 5000 primeiras)" if len(linhas) - 1 > 5000 else "")
              + f", {len(cabecalho)} colunas ({', '.join(cabecalho[:12])})."]

    for i, col in enumerate(cabecalho[:12]):
        valores = [l[i].strip() for l in dados if i < len(l) and l[i].strip()]
        if not valores:
            continue
        numeros = [n for n in (_para_numero(v) for v in valores) if n is not None]
        if len(numeros) >= len(valores) * 0.8:  # coluna numerica
            partes.append(
                f"{col}: min {min(numeros):g}, max {max(numeros):g}, "
                f"media {sum(numeros) / len(numeros):.2f}, "
                f"soma {sum(numeros):g}")
        else:
            top = Counter(valores).most_common(3)
            resumo = ", ".join(f"{v} ({n}x)" for v, n in top)
            partes.append(f"{col}: {len(set(valores))} valores distintos; "
                          f"mais comuns: {resumo}")
    saida = "\n".join(partes)
    return saida[:2000]
